"""
Document assembler: extracts text from Russian .docx and builds clean translated .docx.
Uses python-docx for both reading and writing.
"""

import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def extract_text_from_docx(filepath: str | Path) -> str:
    """
    Extract plain text from a .docx file, preserving paragraph breaks.
    Returns text with single newlines between paragraphs.
    """
    doc = Document(filepath)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        paragraphs.append(text)

    # Join with single newlines; consecutive empty paragraphs become blank lines
    full_text = "\n".join(paragraphs)

    # Collapse runs of 3+ newlines into 2 (one blank line between paragraphs)
    full_text = re.sub(r"\n{3,}", "\n\n", full_text)

    logger.info(
        f"Extracted {len(full_text)} characters, "
        f"{len(full_text.split())} words from {filepath}"
    )
    return full_text


def build_translated_docx(
    translated_text: str,
    output_path: str | Path,
    source_filename: str = "",
) -> Path:
    """
    Build a clean .docx from translated text.

    Formatting rules:
    - No empty lines between paragraphs — Word paragraph spacing handles the gap
    - When copy-pasted to a forum, each paragraph becomes its own block with
      natural spacing, so no extra blank lines are needed in the source
    - Timestamps on their own line, bold
    - Clean, readable font (Calibri 11pt)
    """
    output_path = Path(output_path)
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Paragraph spacing — Word handles visual gaps between paragraphs.
    # No empty lines needed in the text itself.
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.15

    # Split into paragraphs and add to document, skipping empty lines
    lines = translated_text.split("\n")
    timestamp_pattern = re.compile(r"^\d{2}[.:]\d{2}[.:]\d{2,3}$")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue  # Skip empty lines — paragraph spacing handles the gap

        para = doc.add_paragraph()

        if timestamp_pattern.match(stripped):
            run = para.add_run(stripped)
            run.bold = True
        else:
            para.add_run(stripped)

    doc.save(output_path)
    logger.info(f"Saved translated document to {output_path}")
    return output_path


def generate_output_filename(source_filename: str) -> str:
    """
    Generate an English output filename from the Russian source filename.
    Example: "ПУ_102_Оральный_вектор_2_09_03_2026_Часть_1.docx"
           → "PU_102_Oral_Vector_2_09_03_2026_Part_1.docx"

    Falls back to adding "_EN" suffix if pattern doesn't match.
    """
    name = Path(source_filename).stem
    ext = Path(source_filename).suffix or ".docx"

    # Try common substitutions
    replacements = {
        "ПУ": "PU",
        "Оральный_вектор": "Oral_Vector",
        "Оральный вектор": "Oral_Vector",
        "Звуковой_вектор": "Auditory_Vector",
        "Звуковой вектор": "Auditory_Vector",
        "Зрительный_вектор": "Visual_Vector",
        "Зрительный вектор": "Visual_Vector",
        "Кожный_вектор": "Dermal_Vector",
        "Кожный вектор": "Dermal_Vector",
        "Анальный_вектор": "Anal_Vector",
        "Анальный вектор": "Anal_Vector",
        "Уретральный_вектор": "Urethral_Vector",
        "Уретральный вектор": "Urethral_Vector",
        "Мышечный_вектор": "Muscular_Vector",
        "Мышечный вектор": "Muscular_Vector",
        "Обонятельный_вектор": "Olfactory_Vector",
        "Обонятельный вектор": "Olfactory_Vector",
        "Часть": "Part",
        "часть": "Part",
    }

    result = name
    for ru, en in replacements.items():
        result = result.replace(ru, en)

    # If the name still contains Cyrillic, just append _EN
    if re.search(r"[а-яА-ЯёЁ]", result):
        result = name + "_EN"

    return result + ext
